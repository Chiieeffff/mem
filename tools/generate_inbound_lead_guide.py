from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
C_TEAL  = RGBColor(0x00, 0x8B, 0x8B)
C_MINT  = RGBColor(0xE8, 0xF8, 0xF5)
C_LGREY = RGBColor(0xF4, 0xF4, 0xF4)
C_AMBER = RGBColor(0xFF, 0xF3, 0xCD)
C_RED   = RGBColor(0xC0, 0x39, 0x2B)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x1A, 0x1A, 0x1A)

FONT = "Calibri"


def rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    tcPr.append(shd)


def run(para, text, bold=False, italic=False, size=11, color=C_BLACK):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = FONT
    return r


def spacer(doc, pts=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)


def section_header(doc, number, title, subtitle=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2) if subtitle else Pt(7)
    run(p, f"{number}  ", bold=True, size=9, color=C_TEAL)
    run(p, title.upper(), bold=True, size=13, color=C_WHITE)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(7)
        run(p2, subtitle, italic=True, size=10, color=RGBColor(0xAA, 0xCC, 0xCC))
    spacer(doc)


def shaded_block(doc, items, bg=C_MINT, bullet=True, bold_prefix=False):
    """Items can be strings or (bold_label, rest_of_text) tuples."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    first = True
    for item in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if bullet:
            run(p, "•  ", bold=True, size=10, color=C_TEAL)
        if isinstance(item, tuple):
            run(p, item[0], bold=True, size=10.5, color=C_NAVY)
            run(p, "  —  " + item[1], size=10.5, color=C_BLACK)
        else:
            run(p, item, size=10.5)
    spacer(doc)


def sub_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run(p, text, bold=True, size=11, color=C_TEAL)


def script_block(doc, speaker, text):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(0.4)
    tbl.columns[1].width = Cm(15.6)
    set_cell_bg(tbl.rows[0].cells[0], C_TEAL)
    content = tbl.rows[0].cells[1]
    set_cell_bg(content, C_MINT)
    p = content.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    run(p, speaker + "  ", bold=True, size=10, color=C_TEAL)
    run(p, text, italic=True, size=10.5)
    spacer(doc)


def two_col_signals(doc, good, bad):
    tbl = doc.add_table(rows=1 + max(len(good), len(bad)), cols=2)
    tbl.style = "Table Grid"

    for i, (label, bg) in enumerate([("✅  GOOD SIGNALS", C_TEAL), ("🚩  RED FLAGS", C_RED)]):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, bg)
        ph = c.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(5)
        ph.paragraph_format.space_after = Pt(5)
        run(ph, label, bold=True, size=10, color=C_WHITE)

    bgs = [C_MINT, C_AMBER]
    colors = [C_BLACK, C_RED]
    for ri in range(1, len(tbl.rows)):
        for ci, items in enumerate([good, bad]):
            c = tbl.rows[ri].cells[ci]
            set_cell_bg(c, bgs[ci])
            cp = c.paragraphs[0]
            cp.paragraph_format.left_indent = Inches(0.05)
            cp.paragraph_format.space_before = Pt(3)
            cp.paragraph_format.space_after = Pt(3)
            text = items[ri - 1] if ri - 1 < len(items) else ""
            run(cp, ("•  " if text else "") + text, size=10, color=colors[ci])

    spacer(doc)


def objections_table(doc, rows_data):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=2)
    tbl.style = "Table Grid"

    for i, h in enumerate(["OBJECTION", "YOUR RESPONSE"]):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, C_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run(p, h, bold=True, size=10, color=C_WHITE)

    for ri, (obj, resp) in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        oc = row.cells[0]
        set_cell_bg(oc, C_AMBER)
        op = oc.paragraphs[0]
        op.paragraph_format.space_before = Pt(5)
        op.paragraph_format.space_after = Pt(5)
        op.paragraph_format.left_indent = Inches(0.08)
        run(op, obj, bold=True, italic=True, size=10, color=C_RED)
        rc = row.cells[1]
        set_cell_bg(rc, C_LGREY)
        rp = rc.paragraphs[0]
        rp.paragraph_format.space_before = Pt(5)
        rp.paragraph_format.space_after = Pt(5)
        rp.paragraph_format.left_indent = Inches(0.08)
        run(rp, resp, size=10)

    spacer(doc)


def close_options_table(doc, options):
    tbl = doc.add_table(rows=1 + len(options), cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(4)
    tbl.columns[1].width = Cm(12)

    for i, h in enumerate(["SCENARIO", "WHAT TO SAY"]):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, C_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run(p, h, bold=True, size=10, color=C_WHITE)

    bgs = [C_MINT, C_LGREY, C_MINT]
    for ri, (label, text) in enumerate(options):
        row = tbl.rows[ri + 1]
        lc = row.cells[0]
        set_cell_bg(lc, C_TEAL)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after = Pt(8)
        run(lp, label, bold=True, size=10, color=C_WHITE)
        tc = row.cells[1]
        set_cell_bg(tc, bgs[ri])
        tp = tc.paragraphs[0]
        tp.paragraph_format.space_before = Pt(5)
        tp.paragraph_format.space_after = Pt(5)
        tp.paragraph_format.left_indent = Inches(0.08)
        run(tp, text, italic=True, size=10)

    spacer(doc)


def summary_table(doc, fields):
    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(5.5)
    tbl.columns[1].width = Cm(10.5)

    for ri, label in enumerate(fields):
        lc = tbl.rows[ri].cells[0]
        set_cell_bg(lc, C_NAVY)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(5)
        lp.paragraph_format.space_after = Pt(5)
        lp.paragraph_format.left_indent = Inches(0.08)
        run(lp, label, bold=True, size=10, color=C_WHITE)
        vc = tbl.rows[ri].cells[1]
        set_cell_bg(vc, C_LGREY if ri % 2 == 0 else C_MINT)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(5)
        vp.paragraph_format.space_after = Pt(5)

    spacer(doc)


# ── Build ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Title ────────────────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_NAVY)

    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(16)
    run(tp, "GATEWAY CRYPTO", bold=True, size=24, color=C_WHITE)

    tp2 = cell.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(tp2, "Inbound Lead Call Guide", size=13, color=RGBColor(0xAA, 0xCC, 0xCC))

    tp3 = cell.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp3.paragraph_format.space_after = Pt(16)
    run(tp3, "For prospects who filled out the website form  ·  Target: 20–30 min",
        italic=True, size=10, color=RGBColor(0x77, 0x99, 0x99))

    spacer(doc, 6)

    # ── 01 Before the Call ───────────────────────────────────────────────────
    section_header(doc, "01", "Before the Call", "What you know — and what you don't")

    sub_header(doc, "What you have from the form:")
    shaded_block(doc, [
        "Their name and company",
        "Their website — look it up before the call. 2 minutes of research, pays off every time.",
        "Their phone number (they want to be contacted)",
    ], bg=C_MINT)

    sub_header(doc, "What you don't know yet:")
    shaded_block(doc, [
        "Their industry and business model",
        "Their role — are they the decision-maker or an evaluator?",
        "Their volume / turnover",
        "Whether they already process crypto or are starting fresh",
        "What specifically triggered them to fill out the form",
        "Their geos",
    ], bg=C_AMBER)

    # ── 02 Opener ────────────────────────────────────────────────────────────
    section_header(doc, "02", "Opener", "First 2 minutes — they came to you, stay curious")
    script_block(doc, "YOU →",
        '"Hey [Name], it\'s [your name] from Gateway Crypto — I saw you filled out the form on our '
        'website, so I wanted to jump on a quick call. I\'ll keep it short. My goal is just to understand '
        'a bit about what you\'re building, then give you a quick picture of what Gateway does, and we\'ll '
        'figure out if there\'s actually a fit here. Sound good?"')
    shaded_block(doc,
        ["They came to you — so don't over-pitch the opener. Set the frame as a conversation, not a demo."],
        bg=RGBColor(0xFF, 0xF9, 0xE6), bullet=False)

    # ── 03 Qualification + Discovery ─────────────────────────────────────────
    section_header(doc, "03", "Qualification + Discovery",
                   "Understand who you're talking to and whether they qualify — before you pitch anything")

    sub_header(doc, "Part A — Who is this person?  (B2B qualification)")
    shaded_block(doc, [
        '"What\'s your role at the company — are you on the ops or finance side, or more on the technical side?"',
        '"Is it just you looking into this, or is there a team or co-founder involved in the decision?"',
        '"Who typically signs off on a tool like this at your company?"',
    ], bg=C_MINT)
    shaded_block(doc,
        ["Goal: establish whether you're talking to a decision-maker or an evaluator. If evaluator, ask who the decision-maker is and whether they can be looped in."],
        bg=RGBColor(0xFF, 0xF9, 0xE6), bullet=False)

    sub_header(doc, "Part B — What's the business?  (Discovery)")
    shaded_block(doc, [
        '"What kind of business are you running — gaming, forex, something else entirely?"',
        '"Are you already processing crypto, or is this something you\'re looking to add?"',
        '"What currencies or chains do your users mostly use — USDT, BTC, others?"',
        '"Roughly what monthly processing volume are you at — or expecting to reach?"',
        '"What geos are you focused on — where do most of your users or clients come from?"',
        '"What prompted you to reach out now — is there a specific problem you\'re running into, or more of a timing thing?"',
    ], bg=C_MINT)

    sub_header(doc, "What to listen for:")
    two_col_signals(doc,
        [
            "Understands crypto basics (chains, USDT, gas fees)",
            "Has real or near-term volume — or is gaming with traction",
            "Knows their geos clearly",
            "Has a specific pain point or deadline (urgency)",
            "Is the decision-maker, or has direct access to one",
        ],
        [
            "US retail is primary geo or target market",
            "Non-gaming volume below $75K/month with no clear growth path",
            "\"Just exploring, no timeline\" with no concrete need",
            "Evaluator with no access to decision-maker",
            "Business feels like a cover-up — evasive on what they actually do",
        ]
    )

    # ── 04 The Pitch ─────────────────────────────────────────────────────────
    section_header(doc, "04", "The Pitch",
                   "Use what you just learned — pick the points that fit, don't recite them all")
    shaded_block(doc, [
        ("Network-agnostic USDT/USDC",
         "No cross-chain swap fees, no manual ops, no risk of chain-specific liquidity running dry at 4am. If they understand crypto, this lands immediately."),
        ("Zero merchant churn in 18 months",
         "Nobody who started with Gateway has left for a competitor. In this industry, that's the strongest proof point available."),
        ("24-hour go-live",
         "KYB and integration run in parallel. If they move, they can be live in a day."),
        ("Customizable role management",
         "Merchants replicate their existing security and approval protocols inside the platform. It adapts to them — not the other way around."),
        ("Simplified end-user KYC",
         "First deposit only, then up to $10K without re-verification. Frictionless for players, clean for the operator."),
        ("24/7 integration support + dedicated account manager",
         "Covered at every stage — not just during onboarding."),
        ("Never been hacked · 99.9% uptime",
         "Can't prove security until something goes wrong. Nothing has in 18 months."),
        ("Full-cycle pricing transparency",
         "Most providers quote deposit-only. Gateway prices the full cycle — no surprises when the invoice arrives."),
    ], bg=C_MINT)

    # ── 05 Handling Objections ───────────────────────────────────────────────
    section_header(doc, "05", "Handling Objections")
    objections_table(doc, [
        (
            '"We already have a crypto processor."',
            '"What\'s your biggest frustration with them right now?" — Wait. There\'s always something. Then position Gateway against that specific thing.'
        ),
        (
            '"Sounds expensive."',
            '"What are you paying now, full cycle — deposit, exchange, withdrawal? Most people are quoting just the deposit fee. Let\'s compare properly before deciding."'
        ),
        (
            '"We need to think about it / talk to the team."',
            '"Totally fair. What specifically do you need clarity on? I can get you a one-pager and set up a quick demo — that\'s usually enough to decide either way. What\'s your timeline?"'
        ),
        (
            '"Not sure about the KYB / compliance process."',
            '"It\'s one page. They work with high-risk businesses directly — no need for storefronts or workarounds. I can get the doc to you today."'
        ),
        (
            '"We have users from [grey-area country]."',
            '"Let me confirm that specific case with the team — it\'s usually fine but I\'d rather give you the right answer than guess."'
        ),
        (
            '"We\'re not ready yet / too early stage."',
            '"Understood. What does \'ready\' look like for you? If there\'s a timeline, I\'d rather stay in touch and intro you at the right moment than push it before you\'re set."'
        ),
    ])

    # ── 06 The Close ─────────────────────────────────────────────────────────
    section_header(doc, "06", "The Close", "Close on a next step — not a decision")

    sub_header(doc, "Closing arguments — use 1 or 2, not all of them:")
    shaded_block(doc, [
        "They came to you — that's half the work done. Most operators need convincing that crypto is worth the effort.",
        "24-hour go-live means no long procurement or integration drag.",
        "The first step — KYB pre-approval — costs them nothing and takes about a week.",
        "18 months, zero churn. If it wasn't working, someone would have left by now.",
    ], bg=C_MINT)

    sub_header(doc, "Next step options:")
    close_options_table(doc, [
        (
            "Option A\nThey're warm",
            '"Let me send you the one-pager today and set up a 20-minute demo with the Gateway team. '
            'You\'ll see the product, ask your questions, and you\'ll know quickly if it\'s the right fit. '
            'Does [day] or [day] work?"'
        ),
        (
            "Option B\nNeeds internal buy-in",
            '"Makes sense. I\'ll send you the one-pager and a quick pricing summary — something concrete '
            'to share with your team. Who else needs to be looped in for a decision like this?"'
        ),
        (
            "Option C\nReady to go",
            '"Perfect. I\'ll intro you to the Gateway team on Telegram. You give them your business type, '
            'geos, and what you need. They pre-approve quickly and set up a group — KYB, integration, '
            'and pricing all handled there. You could be live inside a week."'
        ),
    ])

    # ── 07 What to Send After the Call ───────────────────────────────────────
    section_header(doc, "07", "What to Send After the Call")
    shaded_block(doc, [
        "Gateway Crypto one-pager on product features (get from Irwin / Roman)",
        "Pricing walkthrough — send as a message, never as a PDF",
        "Your Telegram handle or the intro group link",
        "Short follow-up note: what you discussed, what the next step is, and by when",
    ])

    # ── 08 Call Summary Template ─────────────────────────────────────────────
    section_header(doc, "08", "Call Summary Template", "Fill in after every call")
    summary_table(doc, [
        "Prospect name",
        "Company / website",
        "Their role",
        "Decision-maker? (Y / N / Who is?)",
        "Industry / business type",
        "Current crypto setup",
        "Pain point / trigger for reaching out",
        "Geos",
        "Monthly volume (approx)",
        "Red flags (if any)",
        "Next step agreed",
        "Follow-up date",
    ])

    out = "/Users/roman/Downloads/Gateway_Crypto_Inbound_Lead_Guide.docx"
    doc.save(out)
    print(f"Saved → {out}")

build()
