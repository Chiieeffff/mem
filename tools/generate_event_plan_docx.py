from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
C_TEAL  = RGBColor(0x00, 0x8B, 0x8B)
C_MINT  = RGBColor(0xE8, 0xF8, 0xF5)
C_LGREY = RGBColor(0xF4, 0xF4, 0xF4)
C_AMBER = RGBColor(0xFF, 0xF3, 0xCD)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
C_GREEN = RGBColor(0x1A, 0x7A, 0x4A)

FONT = "Calibri"


def rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    tcPr.append(shd)


def run(para, text, bold=False, italic=False, size=11, color=C_BLACK):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = FONT
    return r


def section_header(doc, number, title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    run(p, f"{number}  ", bold=True, size=9, color=C_TEAL)
    run(p, title.upper(), bold=True, size=13, color=C_WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet_block(doc, items, bg=C_MINT, icon="•"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    first = True
    for item in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            run(p, icon + "  ", bold=True, size=10, color=C_TEAL)
            run(p, item[0], bold=True, size=10.5, color=C_BLACK)
            run(p, "  " + item[1], size=10, color=C_BLACK)
        else:
            run(p, icon + "  ", bold=True, size=10, color=C_TEAL)
            run(p, item, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(doc, text, bg=C_AMBER):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(0.4)
    tbl.columns[1].width = Cm(15.6)
    set_cell_bg(tbl.rows[0].cells[0], C_TEAL)
    set_cell_bg(tbl.rows[0].cells[1], bg)
    p = tbl.rows[0].cells[1].paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    run(p, text, italic=True, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def meeting_card(doc, idx, company, type_, contact, website, co_li, contact_li, angle):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(1.4)
    tbl.columns[1].width = Cm(14.6)

    nc = tbl.rows[0].cells[0]
    set_cell_bg(nc, C_TEAL)
    np_ = nc.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_before = Pt(10)
    run(np_, str(idx), bold=True, size=16, color=C_WHITE)

    tc = tbl.rows[0].cells[1]
    set_cell_bg(tc, C_MINT)

    tp = tc.paragraphs[0]
    tp.paragraph_format.space_before = Pt(5)
    tp.paragraph_format.left_indent = Inches(0.1)
    run(tp, company, bold=True, size=12, color=C_NAVY)
    run(tp, f"  |  {type_}", size=10, color=C_TEAL)

    cp = tc.add_paragraph()
    cp.paragraph_format.left_indent = Inches(0.1)
    run(cp, "Contact: ", bold=True, size=10, color=C_BLACK)
    run(cp, contact, size=10)

    wp = tc.add_paragraph()
    wp.paragraph_format.left_indent = Inches(0.1)
    run(wp, "Website: ", bold=True, size=10, color=C_BLACK)
    run(wp, website, size=10, color=C_TEAL)

    lp = tc.add_paragraph()
    lp.paragraph_format.left_indent = Inches(0.1)
    run(lp, "LinkedIn: ", bold=True, size=10, color=C_BLACK)
    run(lp, co_li, size=10, color=C_TEAL)
    if contact_li:
        run(lp, "  |  ", size=10, color=C_BLACK)
        run(lp, contact_li, size=10, color=C_TEAL)

    ap = tc.add_paragraph()
    ap.paragraph_format.space_after = Pt(6)
    ap.paragraph_format.left_indent = Inches(0.1)
    run(ap, "Angle: ", bold=True, size=10, color=C_TEAL)
    run(ap, angle, italic=True, size=10, color=C_BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def pipeline_table(doc, rows):
    headers = ["COMPANY", "TYPE", "CONTACT", "STAND", "ACTION"]
    col_widths = [Cm(3.8), Cm(3.2), Cm(3.2), Cm(1.8), Cm(4.0)]
    tbl = doc.add_table(rows=1 + len(rows), cols=5)
    tbl.style = "Table Grid"
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = w

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, C_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run(p, h, bold=True, size=9, color=C_WHITE)

    bgs = [C_MINT, C_LGREY]
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            set_cell_bg(c, bgs[ri % 2])
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.06)
            bold = ci == 0
            color = C_NAVY if ci == 0 else C_TEAL if ci == 3 and val != "—" else C_BLACK
            run(p, val, bold=bold, size=9, color=color)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def schedule_table(doc, rows):
    headers = ["DATE / TIME", "ACTIVITY", "PRIORITY"]
    col_widths = [Cm(3.5), Cm(9.5), Cm(3.0)]
    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    tbl.style = "Table Grid"
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = w

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, C_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run(p, h, bold=True, size=9, color=C_WHITE)

    bgs = [C_MINT, C_LGREY]
    priority_colors = {"HIGH": C_GREEN, "MEDIUM": C_TEAL, "TBC": RGBColor(0x88, 0x88, 0x88)}
    for ri, (date, activity, priority) in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate([date, activity, priority]):
            c = row.cells[ci]
            set_cell_bg(c, bgs[ri % 2])
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.06)
            color = priority_colors.get(val, C_BLACK) if ci == 2 else C_BLACK
            run(p, val, bold=(ci == 2), size=9, color=color)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Title block
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_NAVY)

    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(16)
    run(tp, "next.io 2026", bold=True, size=26, color=C_WHITE)

    tp2 = cell.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(tp2, "Event Strategy Plan  —  GatewayCrypto", size=13, color=RGBColor(0xAA, 0xCC, 0xCC))

    tp3 = cell.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp3.paragraph_format.space_after = Pt(16)
    run(tp3, "Prepared by Roman  |  Pre-registration: May 26, 2026",
        italic=True, size=10, color=RGBColor(0x77, 0x99, 0x99))

    doc.add_paragraph()

    # ── 01 Overview ─────────────────────────────────────────────────────────────
    section_header(doc, "01", "Overview")
    bullet_block(doc, [
        ("Event:", "next.io — iGaming Expo + Collective"),
        ("Goal:", "Generate qualified pipeline from iGaming platforms, operators, and payment partners by securing intro meetings, qualifying fit, and locking concrete next steps before leaving."),
        ("GatewayCrypto:", "Crypto payment processing for iGaming — 300+ assets, unified USDT/USDC across chains, zero chargebacks, EUR off-ramp via SEPA/SWIFT, live within 24 hours."),
        ("Pipeline snapshot:", "13 companies engaged  |  6 confirmed intro meetings  |  7 in-process"),
    ])

    # ── 02 Objectives & Metrics ──────────────────────────────────────────────────
    section_header(doc, "02", "Objectives & Success Metrics")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    run(p, "Pre-Event Targets", bold=True, size=11, color=C_TEAL)
    bullet_block(doc, [
        ("Confirmed meetings booked:", "6/6  ✓"),
        ("In-process converted before event day:", "Target 4/7 (57%)"),
        ("Stand walk-bys planned:", "AMG Platform (A4)  ·  Gaming Ent (L18)"),
    ], bg=C_LGREY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(3)
    run(p2, "At-Event Targets", bold=True, size=11, color=C_TEAL)
    bullet_block(doc, [
        "Rate each conversation: Hot / Warm / Cold",
        "Leave every confirmed meeting with 1 concrete next step (demo, intro, or trial)",
        "Collect 1 specific objection or blocker per cold/warm conversation",
    ], bg=C_LGREY)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(3)
    run(p3, "Post-Event Targets (30-day horizon)", bold=True, size=11, color=C_TEAL)
    bullet_block(doc, [
        ("Follow-up sent within 48h:", "100% of met contacts"),
        ("Demo or discovery call scheduled:", "4+ within 2 weeks"),
        ("Proposal / commercial discussion opened:", "2+ within 30 days"),
        ("New qualified opportunities in tracker:", "6+"),
    ], bg=C_LGREY)

    # ── 03 Confirmed Meetings ────────────────────────────────────────────────────
    section_header(doc, "03", "Confirmed Intro Meetings (6)")

    meetings = [
        (1, "Plaee", "Platform", "Libi Milshtein",
         "https://www.plaee.com/", "linkedin.com/company/plaee/", "linkedin.com/in/libi-milshtein/",
         "Platform-level integration — every operator on Plaee inherits crypto automatically. 300+ assets, unified USDT/USDC, zero chargebacks across the board."),
        (2, "Morefin", "PSP Orchestrator", "Miroslav Naydenov",
         "https://www.morefin.com/", "linkedin.com/company/morefin/about/", "linkedin.com/in/mironaydenov/",
         "GWC as the crypto rail inside their orchestration stack. They route card/APM — we handle crypto. Additive, no overlap, one integration expands their full merchant base."),
        (3, "BigBux", "Operator / Ruffle", "Viktor Atanasovski",
         "https://bigbux.io/", "linkedin.com/company/bigbux/about/", "linkedin.com/in/viktoratanasovski/",
         "Direct integration for their player base. Ruffle mechanics attract crypto-native users. Crypto deposits are irreversible — zero chargeback exposure."),
        (4, "BetHero", "Sportsbook", "Mike Forslund",
         "https://bethero.gg/", "linkedin.com/company/hellobethero/", "linkedin.com/in/mikeforslund/",
         "Sports bettors skew crypto-native. Instant settlement, zero chargebacks, 300+ assets. Direct integration — live within 24 hours if they're responsive."),
        (5, "Start2Pay", "PSP", "Tanya Seleznova",
         "https://start2pay.com/", "", "linkedin.com/in/tanya-seleznova-13a04a212/",
         "Partnership model — they cover fiat rails, GWC covers the crypto leg. EUR off-ramp via SEPA/SWIFT bridges the fiat-crypto gap for shared merchants. No competitive overlap."),
        (6, "PayAdmit", "PSP", "Heorhii Kuchuk",
         "https://payadmit.com/", "linkedin.com/company/payadmit/", "linkedin.com/in/heorhii-kuchuk-ab5561247/",
         "Same partnership model as Start2Pay — GWC as the crypto rail they don't have natively. Referral or white-label structure. No competition on the fiat side."),
    ]
    for m in meetings:
        meeting_card(doc, *m)

    # ── 04 In-Process Pipeline ───────────────────────────────────────────────────
    section_header(doc, "04", "In-Process Pipeline (7)")

    callout(doc,
        "AMG (A4) and Gaming Ent (L18) have confirmed stand numbers — approach on the floor if pre-event intro is not locked.",
        bg=C_AMBER)

    pipeline_table(doc, [
        ("AMG Platform",     "Platform",             "Mark Abdilla",           "A4",   "Walk-by confirmed"),
        ("Gaming Ent",       "Platform",             "Boyko Boev",             "L18",  "Walk-by confirmed"),
        ("Kanggiten",        "Platform",             "Sergey Shibkih",         "—",    "Chase pre-event"),
        ("Slikair",          "Payment Orchestrator", "Tzach Toporek",          "—",    "Chase pre-event"),
        ("Comtrade Gaming",  "Platform",             "S. Valentine / A. Gornjec", "—", "Chase pre-event — 2 contacts"),
        ("Bejoynd",          "Platform",             "Fredrik Cedell",         "—",    "Chase pre-event"),
        ("Ace Systems",      "Platform",             "William Lövqvist",       "—",    "Chase pre-event"),
    ])

    # ── 05 Strategy by Company Type ──────────────────────────────────────────────
    section_header(doc, "05", "Targeting Strategy by Company Type")
    bullet_block(doc, [
        ("Platforms (8 companies):", "Platform-level integration pitch — operators inherit crypto by default. One deal unlocks the entire operator base. Highest pipeline leverage at this event."),
        ("PSPs & Orchestrators (4 companies):", "Partnership model — GWC as the crypto rail in their stack. They keep their fiat business. We add the crypto leg. No competitive overlap."),
        ("Operators & Sportsbook (2 companies):", "Direct integration — player-facing crypto deposits and withdrawals. Zero chargebacks, irreversible transactions, 24h go-live."),
    ])

    # ── 06 Event Schedule ────────────────────────────────────────────────────────
    section_header(doc, "06", "Event Schedule")
    schedule_table(doc, [
        ("May 26",   "Pre-registration event",                           "HIGH"),
        ("TBC",      "Rooftop side event #1",                            "HIGH"),
        ("TBC",      "Rooftop side event #2",                            "HIGH"),
        ("TBC",      "Networking event",                                  "MEDIUM"),
        ("Expo days", "6 confirmed meeting slots",                        "HIGH"),
        ("Expo days", "Walk-bys: AMG (Stand A4), Gaming Ent (Stand L18)", "HIGH"),
    ])
    callout(doc,
        "Update with specific times as the schedule is confirmed. Prioritise morning slots for confirmed meetings — energy and attention are highest early in the day at expos.",
        bg=C_MINT)

    # ── 07 Post-Event Action Plan ─────────────────────────────────────────────────
    section_header(doc, "07", "Post-Event Action Plan")
    bullet_block(doc, [
        ("T+48h:", "Send follow-up LinkedIn message to all 6 confirmed meeting contacts  (reply_drafting.md workflow)"),
        ("T+48h:", "Send outreach to in-process contacts not converted pre-event  (linkedin_outreach.md workflow)"),
        ("T+2 weeks:", "Book demos / discovery calls for all Hot-rated conversations"),
        ("T+2 weeks:", "Update outreach tracker with outcome for every company touched  (outreach_tracker.md)"),
        ("T+30 days:", "Open proposal or commercial discussion for 2+ qualified leads"),
        ("T+30 days:", "Review in-process pipeline — move or close stalled conversations  (telegram_followup_batch.md)"),
    ], icon="→")

    out = "/Users/roman/Downloads/nextio_2026_Event_Strategy_Plan.docx"
    doc.save(out)
    print(f"Saved → {out}")


build()
