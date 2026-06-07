from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
C_TEAL  = RGBColor(0x00, 0x8B, 0x8B)
C_MINT  = RGBColor(0xE8, 0xF8, 0xF5)
C_LGREY = RGBColor(0xF4, 0xF4, 0xF4)
C_AMBER = RGBColor(0xFF, 0xF3, 0xCD)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
C_GREEN = RGBColor(0x1A, 0x7A, 0x4A)

FONT = "Calibri"


def rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    tcPr.append(shd)


def run(para, text, bold=False, italic=False, size=11, color=C_BLACK):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = FONT
    return r


def section_header(doc, number, title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    run(p, f"{number}  ", bold=True, size=9, color=C_TEAL)
    run(p, title.upper(), bold=True, size=13, color=C_WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet_block(doc, items, bg=C_MINT, icon="•"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    first = True
    for item in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            run(p, icon + "  ", bold=True, size=10, color=C_TEAL)
            run(p, item[0], bold=True, size=10.5, color=C_BLACK)
            run(p, "  " + item[1], size=10, color=C_BLACK)
        else:
            run(p, icon + "  ", bold=True, size=10, color=C_TEAL)
            run(p, item, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def numbered_card(doc, number, title, body):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(1.4)
    tbl.columns[1].width = Cm(14.6)

    nc = tbl.rows[0].cells[0]
    set_cell_bg(nc, C_TEAL)
    np_ = nc.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_before = Pt(10)
    run(np_, number, bold=True, size=16, color=C_WHITE)

    tc = tbl.rows[0].cells[1]
    set_cell_bg(tc, C_MINT)
    tp = tc.paragraphs[0]
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.left_indent = Inches(0.1)
    run(tp, title, bold=True, size=11, color=C_NAVY)

    if body:
        bp = tc.add_paragraph()
        bp.paragraph_format.space_after = Pt(6)
        bp.paragraph_format.left_indent = Inches(0.1)
        run(bp, body, size=10, color=C_BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def two_col_table(doc, headers, rows, col_widths=(5.5, 10.5), header_bg=C_NAVY, alt_bg=C_LGREY):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(col_widths[0])
    tbl.columns[1].width = Cm(col_widths[1])

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run(p, h, bold=True, size=10, color=C_WHITE)

    bgs = [C_MINT, alt_bg]
    for ri, (left, right) in enumerate(rows):
        row = tbl.rows[ri + 1]
        lc = row.cells[0]
        set_cell_bg(lc, bgs[ri % 2])
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(5)
        lp.paragraph_format.space_after = Pt(5)
        lp.paragraph_format.left_indent = Inches(0.08)
        run(lp, left, bold=True, size=10, color=C_NAVY)

        rc = row.cells[1]
        set_cell_bg(rc, bgs[ri % 2])
        rp = rc.paragraphs[0]
        rp.paragraph_format.space_before = Pt(5)
        rp.paragraph_format.space_after = Pt(5)
        rp.paragraph_format.left_indent = Inches(0.08)
        run(rp, right, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def pricing_table(doc, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(7)
    tbl.columns[1].width = Cm(9)

    for i, h in enumerate(["FEE TYPE", "DETAIL"]):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, C_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run(p, h, bold=True, size=10, color=C_WHITE)

    row_bgs = [C_MINT, C_LGREY]
    for ri, (fee, detail) in enumerate(rows):
        row = tbl.rows[ri + 1]
        fc = row.cells[0]
        set_cell_bg(fc, row_bgs[ri % 2])
        fp = fc.paragraphs[0]
        fp.paragraph_format.space_before = Pt(5)
        fp.paragraph_format.space_after = Pt(5)
        fp.paragraph_format.left_indent = Inches(0.08)
        run(fp, fee, bold=True, size=10, color=C_TEAL)

        dc = row.cells[1]
        set_cell_bg(dc, row_bgs[ri % 2])
        dp = dc.paragraphs[0]
        dp.paragraph_format.space_before = Pt(5)
        dp.paragraph_format.space_after = Pt(5)
        dp.paragraph_format.left_indent = Inches(0.08)
        run(dp, detail, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(doc, text, bg=C_AMBER):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(0.4)
    tbl.columns[1].width = Cm(15.6)
    set_cell_bg(tbl.rows[0].cells[0], C_TEAL)
    set_cell_bg(tbl.rows[0].cells[1], bg)
    p = tbl.rows[0].cells[1].paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    run(p, text, italic=True, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Build ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Title block
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_NAVY)

    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(16)
    run(tp, "GATEWAY CRYPTO", bold=True, size=24, color=C_WHITE)

    tp2 = cell.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(tp2, "B2B Sales & Partnership Brief", size=13, color=RGBColor(0xAA, 0xCC, 0xCC))

    tp3 = cell.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp3.paragraph_format.space_after = Pt(16)
    run(tp3, "Crypto Payment Processing for High-Risk Businesses",
        italic=True, size=10, color=RGBColor(0x77, 0x99, 0x99))

    doc.add_paragraph()

    # ── 01 What It Is ────────────────────────────────────────────────────────
    section_header(doc, "01", "What It Is")
    bullet_block(doc, [
        ("C2B — Consumer to Business:", "Merchants collecting payments from end users and making payouts. The main use case."),
        ("B2B — Corporate:", "On-ramp/off-ramp, manual large transactions, invoicing, exchange capabilities."),
    ])

    # ── 02 Key Differentiators ───────────────────────────────────────────────
    section_header(doc, "02", "Key Differentiators — Your Sales Arguments")

    numbered_card(doc, "1", "Network-agnostic USDT/USDC  —  the killer pitch",
        "Competitors charge for cross-chain swaps. Gateway doesn't. No fees, no manual ops, no risk "
        "of running dry on a specific chain at 4am when players can't withdraw. If the prospect "
        "understands crypto at all, this lands hard.")

    numbered_card(doc, "2", "Fully customizable role management",
        "Merchants mirror their existing security protocols — e.g. \"any transaction above $X requires "
        "owner approval.\" The processor adapts to them, not the other way around.")

    numbered_card(doc, "3", "Track record",
        "Never hacked  ·  99.9% uptime  ·  Zero merchant churn in 18 months — no one has left for a competitor.")

    numbered_card(doc, "4", "Speed to launch",
        "Up and running within 24 hours if the merchant is responsive. KYB and integration run in parallel.")

    numbered_card(doc, "5", "Full client support",
        "Dedicated account manager + 24/7 integration support throughout the entire client lifecycle.")

    numbered_card(doc, "6", "Simplified KYC for end users",
        "First deposit: simplified KYC. After that: transact freely up to $10K without additional KYC. "
        "Frictionless for players — clean for the operator.")

    # ── 03 Target Clients ────────────────────────────────────────────────────
    section_header(doc, "03", "Target Clients")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run(p, "Direct Merchants", bold=True, size=11, color=C_TEAL)

    two_col_table(doc,
        ["SEGMENT", "NOTES"],
        [
            ("iGaming / Sports Betting", "High volume. No minimum — take all comers, it's a numbers game."),
            ("Licensed Forex", "Lower volume risk, higher quality leads. Requires substance — not cover-up ops."),
            ("High-risk service businesses", "Marketing agencies, lead gen, VPN-adjacent businesses blocked from traditional banking rails."),
            ("Adult", "Growing traction. Not a core focus but accepted."),
        ],
        col_widths=(5.5, 10.5)
    )

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(4)
    run(p2, "Channel Partners — The Partnership Play", bold=True, size=11, color=C_TEAL)

    bullet_block(doc, [
        "iGaming platforms that white-label or aggregate merchants.",
        "If Gateway integrates at platform level, all sub-merchants get access automatically.",
        "This is the B2B partnership angle — one deal unlocks many merchants.",
    ])

    # ── 04 Geo & Compliance ──────────────────────────────────────────────────
    section_header(doc, "04", "Geo & Compliance Rules")

    two_col_table(doc,
        ["GEO / CASE", "RULING"],
        [
            ("USA", "No. No US-registered businesses, no US retail-facing merchants, no US transactions."),
            ("China", "No."),
            ("OFAC-sanctioned countries", "No — but edge cases handled practically. Dual-passport holders, Curaçao entities are usually fine."),
            ("High-risk businesses presenting as-is", "Fine. A storefront in crypto actually raises red flags — not lowers them."),
            ("Russian passport + EU entity", "Usually fine. Flag it and confirm with the team rather than killing the deal."),
        ],
        col_widths=(5.5, 10.5)
    )

    # ── 05 Volume Requirements ───────────────────────────────────────────────
    section_header(doc, "05", "Volume Requirements")
    two_col_table(doc,
        ["MERCHANT TYPE", "MINIMUM"],
        [
            ("iGaming / Betting", "None — take all, sort later."),
            ("Forex", "None — quality almost always converts."),
            ("Services / Adult / Other", "~$75K–$100K/month. Below that: not worth the support overhead."),
        ],
        col_widths=(6, 10)
    )

    # ── 06 Pricing ───────────────────────────────────────────────────────────
    section_header(doc, "06", "Pricing Structure")

    callout(doc,
        "Don't share a PDF — it circulates in the market. Send pricing as a message or walk through it verbally.",
        bg=C_AMBER)

    pricing_table(doc, [
        ("Deposit — no exchange",
         "Base fee. Simple top-up via API or invoice."),
        ("Deposit — with auto-exchange",
         "Slightly higher. User pays in altcoin; merchant accumulates USDT automatically."),
        ("Manual internal exchange",
         "Merchant swaps BTC → USDT inside their dashboard."),
        ("Withdrawal",
         "Free — only network fee. Fees are collected on the deposit side."),
        ("Withdrawal — with exchange",
         "Same rate as deposit. User withdraws in a different currency than deposited."),
        ("C2B on-ramp",
         "Cost passed to end user. Net-zero for the merchant — same as a classic deposit."),
        ("B2B on-ramp / off-ramp",
         "Negotiated per client depending on needs."),
    ])

    callout(doc,
        "Key framing: when a prospect says \"I'm paying 0.4%\" — they're almost certainly quoting "
        "deposit-only. Push them to compare full-cycle cost the same way they'd compare fiat processors.",
        bg=C_MINT)

    # ── 07 Partnership Process ───────────────────────────────────────────────
    section_header(doc, "07", "Partnership Process")

    bullet_block(doc, [
        "Drop the lead into a shared Telegram channel: business type, needs, geos.",
        "Gateway pre-approves in chat — fast filter, no wasted time for either side.",
        "They create a group with your client and handle KYB + integration end-to-end.",
        "You can stay in the chat for transparency or step out — your call.",
    ], icon="→")

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(4)
    run(p3, "Revenue Share Options", bold=True, size=11, color=C_TEAL)

    bullet_block(doc, [
        "Gateway sets up an account for you and settles your cut directly in crypto.",
        "Monthly / bi-weekly reconciliation invoice paid in crypto.",
        "Automated referral dashboard coming ~Q3 — real-time fee tracking (not live yet).",
    ])

    # ── 08 What to Share ─────────────────────────────────────────────────────
    section_header(doc, "08", "What You Can Share with Prospects")
    bullet_block(doc, [
        "One-pager on product features — request from Irwin / Roman.",
        "Pricing — send as a message, never as a PDF.",
        "KYB requirements doc — Gateway sends this once a group is created.",
    ])

    out = "/Users/roman/Downloads/Gateway_Crypto_Product_Brief.docx"
    doc.save(out)
    print(f"Saved → {out}")

build()
